from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET
import re, datetime, math, json, statistics
from collections import defaultdict, Counter

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document

TODAY = datetime.date(2026,5,19)
OUT = Path('/Users/openclaw/Downloads/Wickes_Solar_June_Installation_Master_Planner.xlsx')
DNO_XLSX = Path('/Users/openclaw/.hermes/document_cache/doc_3d786406c2c5_Copy_of_G99_Spreadsheet_New.1.xlsx')
DNO_DOCX = Path('/Users/openclaw/.hermes/document_cache/doc_462d87d18b1f_DNO_Ranking.docx')
PIPE_ID = '15wmbRs6FMDOFsoMb6C1JKINwiJ3oLflRSP1NrbIQ-uA'
INST_ID = '1cJNioXl3o_n7zz0Z9SgDzvzSnoKqZbr14tWY7RKdqBI'

# ---------- helpers ----------
def norm(s): return str(s or '').strip()
def norm_id(s): return re.sub(r'[^0-9]', '', str(s or ''))
def parse_date(s):
    s=norm(s)
    if not s: return None
    for fmt in ['%Y-%m-%d','%d/%m/%Y','%d.%m.%Y','%d-%m-%Y','%d/%m/%y']:
        try: return datetime.datetime.strptime(s,fmt).date()
        except ValueError: pass
    try:
        f=float(s)
        if 30000<f<60000:
            return datetime.date(1899,12,30)+datetime.timedelta(days=int(f))
    except Exception: pass
    return None
def fmt_date(d): return d.isoformat() if isinstance(d, datetime.date) else ''
def add_workdays(d, n):
    if not d: d=TODAY
    out=d; added=0
    while added<n:
        out += datetime.timedelta(days=1)
        if out.weekday()<5: added += 1
    return out
def workdays_between(a,b):
    if not a or not b: return None
    step=1 if b>=a else -1; cur=a; count=0
    while cur != b:
        cur += datetime.timedelta(days=step)
        if cur.weekday()<5: count += step
    return count

def clean_postcode(pc): return re.sub(r'\s+','', norm(pc).upper())
def pc_area(pc):
    m=re.match(r'^([A-Z]{1,2})', clean_postcode(pc))
    return m.group(1) if m else ''
def pc_outward(pc):
    pc=clean_postcode(pc)
    return pc[:-3] if len(pc)>3 else pc

# ---------- Google Sheets ----------
def sheets_service():
    token = Path.home()/'.hermes/google_token_wickes_solar.json'
    creds = Credentials.from_authorized_user_file(str(token))
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
    return build('sheets','v4',credentials=creds,cache_discovery=False)

def gvals(svc, sid, tab):
    return svc.spreadsheets().values().get(spreadsheetId=sid, range=f"'{tab}'!A1:ZZ10000", valueRenderOption='FORMATTED_VALUE').execute().get('values',[])
def rows_from_values(vals):
    if not vals: return [], []
    header=[norm(x) for x in vals[0]]
    rows=[]
    for r in vals[1:]:
        if any(norm(x) for x in r):
            rows.append({header[i] if i<len(header) and header[i] else f'col{i+1}': (r[i] if i<len(r) else '') for i in range(len(header))})
    return header, rows

# ---------- xlsx parser ----------
NS={'a':'http://schemas.openxmlformats.org/spreadsheetml/2006/main','r':'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
def col_to_idx(col):
    n=0
    for ch in col: n=n*26+ord(ch)-64
    return n-1
def excel_value(v):
    if v is None: return ''
    s=str(v)
    try:
        f=float(s)
        if 30000 < f < 60000:
            return (datetime.date(1899,12,30)+datetime.timedelta(days=int(f))).isoformat()
    except Exception: pass
    return s

def load_xlsx_sheet(z, name):
    wb=ET.fromstring(z.read('xl/workbook.xml'))
    sheets=wb.findall('a:sheets/a:sheet',NS)
    sh=next(s for s in sheets if s.attrib['name']==name)
    rid=sh.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id']
    rels=ET.fromstring(z.read('xl/_rels/workbook.xml.rels'))
    target=next(rel.attrib['Target'] for rel in rels if rel.attrib['Id']==rid)
    if not target.startswith('xl/'): target='xl/'+target.lstrip('/')
    shared=[]
    if 'xl/sharedStrings.xml' in z.namelist():
        sst=ET.fromstring(z.read('xl/sharedStrings.xml'))
        for si in sst.findall('a:si',NS):
            shared.append(''.join(t.text or '' for t in si.iterfind('.//a:t',NS)))
    root=ET.fromstring(z.read(target)); rows=[]
    for row in root.findall('.//a:sheetData/a:row',NS):
        vals={}
        for c in row.findall('a:c',NS):
            m=re.match(r'([A-Z]+)', c.attrib.get('r',''))
            if not m: continue
            ci=col_to_idx(m.group(1)); t=c.attrib.get('t'); v=c.find('a:v',NS); isel=c.find('a:is',NS)
            val=''
            if t=='s' and v is not None: val=shared[int(v.text)]
            elif t=='inlineStr' and isel is not None: val=''.join(x.text or '' for x in isel.iterfind('.//a:t',NS))
            elif v is not None: val=v.text or ''
            vals[ci]=excel_value(val)
        if vals:
            rows.append([vals.get(i,'') for i in range(max(vals)+1)])
    return rows

def dict_rows_any(rows):
    header_idx=0
    for i,r in enumerate(rows):
        if sum(1 for x in r if norm(x))>=3:
            header_idx=i; break
    header=[norm(x) for x in rows[header_idx]]
    out=[]
    for r in rows[header_idx+1:]:
        if any(norm(x) for x in r):
            d={header[i] if i<len(header) and header[i] else f'col{i+1}': (r[i] if i<len(r) else '') for i in range(max(len(header),len(r)))}
            out.append(d)
    return header,out

# ---------- DNO assumptions ----------
DNO_RULES = {
    'UK Power Networks': {'rank':1,'fast':7,'typical':10,'conservative':14,'risk':'Low','note':'Strong performer; some instant/fast approvals.'},
    'National Grid Electricity Distribution': {'rank':2,'fast':7,'typical':10,'conservative':14,'risk':'Low','note':'Fast-track often 3-14 working days; documentation-sensitive.'},
    'Electricity North West': {'rank':3,'fast':7,'typical':10,'conservative':14,'risk':'Low','note':'Strong turnaround for straightforward applications.'},
    'GTC': {'rank':4,'fast':10,'typical':14,'conservative':45,'risk':'Medium','note':'Simple around 14 working days; complex 45-65.'},
    'Last Mile': {'rank':5,'fast':7,'typical':10,'conservative':14,'risk':'Low/Medium','note':'Generally positive, same-day to 14 working days.'},
    'ESP Utilities Group': {'rank':6,'fast':10,'typical':14,'conservative':45,'risk':'Medium','note':'Standard 10-14; complex 45-65.'},
    'Northern Powergrid': {'rank':7,'fast':14,'typical':21,'conservative':45,'risk':'High','note':'Backlogs; fast-track about 14, otherwise 45-65.'},
    'SP Energy Networks': {'rank':8,'fast':14,'typical':20,'conservative':25,'risk':'High','note':'Moderate with export limits/network constraints.'},
    'Scottish and Southern Electricity Networks': {'rank':9,'fast':20,'typical':30,'conservative':45,'risk':'High','note':'Can be slow; technical requirements around 6kW.'},
    'Energy Assets': {'rank':10,'fast':30,'typical':45,'conservative':65,'risk':'Very High','note':'Extended/inconsistent delays.'},
    'Unknown': {'rank':99,'fast':20,'typical':30,'conservative':45,'risk':'Unknown','note':'Not mapped; verify before confirming.'},
}
# Approx postcode-area mapping; actual DNO tracker/application evidence should override in operations.
UKPN=set('AL BR CB CM CO CR CT DA E EC EN HA IG IP LU ME N NN? NR NW PE RM SE SG SM SS SW TN W WC'.split())
NGED=set('B BA BS CF CV DE DY EX GL HR LE LD LN NG NN NP OX PL SA SN ST SY TA TQ TR WR WS WV'.split())
ENW=set('BB BL CA CH CW FY LA L M OL PR SK WA WN'.split())
NPG=set('BD DH DL DN HG HD HU HX LS NE SR TS WF YO'.split())
SPEN=set('DG EH FK G KA KY ML PA TD LL CH L SY'.split())
SSEN=set('AB BH DD DT GU IV KW PH PO RG SL SO SP ZE'.split())

def dno_from_postcode(pc, notes=''):
    txt=(notes or '').lower()
    if 'npg' in txt or 'northern powergrid' in txt: return 'Northern Powergrid', 'notes keyword'
    if 'nged' in txt or 'national grid' in txt or 'western power' in txt: return 'National Grid Electricity Distribution', 'notes keyword'
    if 'ukpn' in txt or 'uk power' in txt: return 'UK Power Networks', 'notes keyword'
    if 'spen' in txt or 'sp energy' in txt: return 'SP Energy Networks', 'notes keyword'
    if 'ssen' in txt or 'southern electric' in txt or 'scottish and southern' in txt: return 'Scottish and Southern Electricity Networks', 'notes keyword'
    if 'electricity north west' in txt or 'enw' in txt: return 'Electricity North West', 'notes keyword'
    a=pc_area(pc)
    # Known exceptions/priority decisions for common areas in this dataset
    if a in {'NN'}: return 'National Grid Electricity Distribution', 'postcode area'
    if a in {'PE','NR','IP','SG','AL','ME','RM','TN','CO','CM'}: return 'UK Power Networks', 'postcode area'
    if a in {'NE','YO','HG','BD','HD','HX','WF','LS','DN','S','TS','SR','DH','DL'}: return 'Northern Powergrid', 'postcode area'
    if a in {'TR','PL','EX','TQ','TA','BA','BS','SN','OX','B','CV','LE','NG','DE','ST','HR','GL','WR','WV','DY','WS','CF','SA','NP','SY'}: return 'National Grid Electricity Distribution', 'postcode area'
    if a in {'LA','M','BL','BB','PR','FY','OL','SK','WA','WN','CW'}: return 'Electricity North West', 'postcode area'
    if a in {'G','KA','FK','EH','ML','TD','DG','LL','CH'}: return 'SP Energy Networks', 'postcode area'
    if a in {'AB','IV','PH','DD','PO','SO','RG','SL','BH','DT','GU','SP'}: return 'Scottish and Southern Electricity Networks', 'postcode area'
    return 'Unknown', 'unmapped postcode area'

def planning_region(pc, city=''):
    a=pc_area(pc)
    if a in {'NE','DH','SR','TS','DL'}: return 'North East'
    if a in {'LS','WF','BD','HD','HX','HG','YO','HU','DN','S'}: return 'Yorkshire / Humberside'
    if a in {'M','BL','BB','PR','FY','LA','OL','SK','WA','WN','CW','CH','L'}: return 'North West'
    if a in {'LL','SY'}: return 'North Wales'
    if a in {'CF','NP','SA','LD'}: return 'South Wales'
    if a in {'B','CV','DE','DY','LE','NG','NN','ST','WR','WS','WV','HR','LN'}: return 'Midlands'
    if a in {'PE','NR','IP','CB','SG','AL','LU','CO','CM'}: return 'East / East Anglia'
    if a in {'E','EC','N','NW','SE','SW','W','WC','BR','CR','DA','EN','HA','IG','RM','SM','SS','KT','TW','UB'}: return 'London / South East'
    if a in {'ME','CT','TN','BN','RH','GU'}: return 'Kent / Sussex / South East'
    if a in {'OX','SN','GL','BS','BA','TA','EX','TQ','TR','PL','SP','DT'}: return 'South West / West'
    if a in {'PO','SO','RG','SL','BH'}: return 'South Coast / Thames Valley'
    if a in {'G','KA','FK','EH','ML','DG','TD','PA'}: return 'Scotland Central/South'
    if a in {'AB','DD','IV','PH','KW'}: return 'Scotland North'
    return 'Other / Manual'

# Installer allocation map
INSTALLERS = {
    'Renewable Green Energy': {'region':['North Wales','North West'], 'status':'Live / Green', 'risk':'Preferred where North Wales/NW fit'},
    'Simple Solar': {'region':['South West / West'], 'status':'Live / Green', 'risk':'Preferred where South West fit'},
    'Campbell Electrical & Renewables': {'region':['Yorkshire / Humberside'], 'status':'Live / Green', 'risk':'Preferred where Yorkshire fit'},
    'Solec Electrical Services': {'region':['North East','Yorkshire / Humberside','National cluster'], 'status':'Available / strategic', 'risk':'Use nationally only for clustered consecutive runs'},
    'Eco Solve Renewables': {'region':['South West / West'], 'status':'Compliance candidate', 'risk':'Use where local and docs/pricing confirmed'},
    'Alpha Solar': {'region':['London / South East'], 'status':'Compliance candidate', 'risk':'Use London/SE with zone cost awareness'},
    'Green Light Solar': {'region':['London / South East','Kent / Sussex / South East'], 'status':'Compliance candidate', 'risk':'Use Essex/SE where fit'},
    'RPS Electricals': {'region':['Kent / Sussex / South East','London / South East'], 'status':'Compliance candidate', 'risk':'Kent/SE option; docs follow-up noted'},
    'Qwinteq': {'region':['Midlands','East / East Anglia'], 'status':'Compliance candidate', 'risk':'Midlands option; confirm roofer/docs'},
    'Renewable Edtricity': {'region':['East / East Anglia','Midlands'], 'status':'Compliance candidate', 'risk':'Potential capacity; team numbers unclear'},
    'Prolec Energy Solutions': {'region':['South Wales'], 'status':'Compliance candidate', 'risk':'MCS not fully complete per notes'},
    'Smart Home Technical': {'region':['South West / West'], 'status':'Compliance candidate', 'risk':'Waiting on docs'},
    'Smartliving Energy': {'region':['Scotland Central/South','Scotland North'], 'status':'Live/candidate caution', 'risk':'Difficult to deal with; Scotland controlled use only'},
    'CK Solar': {'region':['National cluster'], 'status':'Caution', 'risk':'Unreliable; avoid/last resort only'},
}
def installer_for(region, pc, dno, risk_bucket):
    # Local preferred first
    if region == 'North Wales': return 'Renewable Green Energy', 'Local North Wales live installer'
    if region == 'South Wales': return 'Prolec Energy Solutions', 'South Wales local option; verify MCS/docs before confirming'
    if region == 'Yorkshire / Humberside': return 'Campbell Electrical & Renewables', 'Local/live Yorkshire fit'
    if region == 'North East': return 'Solec Electrical Services', 'North East base; natural local fit'
    if region == 'South West / West': return 'Simple Solar', 'Live South West fit; Eco Solve/Smart Home backup'
    if region == 'London / South East': return 'Alpha Solar', 'London/SE fit; confirm zone/access costs'
    if region == 'Kent / Sussex / South East': return 'RPS Electricals', 'Kent/SE fit; Green Light backup'
    if region == 'Midlands': return 'Qwinteq', 'Midlands fit; Solec only if clustered'
    if region == 'East / East Anglia': return 'Renewable Edtricity', 'East/Bedfordshire/East Anglia option; confirm capacity'
    if region == 'North West': return 'Renewable Green Energy', 'North Wales/NW fit; cluster if travel heavy'
    if region in ['Scotland Central/South','Scotland North']:
        return 'Smartliving Energy', 'Scotland option but difficult; use controlled slots or Solec cluster backup'
    if region == 'South Coast / Thames Valley': return 'Simple Solar', 'South/SW route fit; cluster to reduce travel'
    return 'Solec Electrical Services', 'Manual allocation; use Solec only if clustered'

# ---------- pull data ----------
svc=sheets_service()
pipe_header, orders = rows_from_values(gvals(svc, PIPE_ID, 'Orders details'))
notes_header, notes = rows_from_values(gvals(svc, PIPE_ID, 'Order Notes'))
inst_header, inst_rows = rows_from_values(gvals(svc, INST_ID, 'Move to compliance tracker'))

notes_by_order=defaultdict(list)
for n in notes:
    oid=norm_id(n.get('OrderId'))
    if oid:
        notes_by_order[oid].append(n)
for oid in notes_by_order:
    notes_by_order[oid].sort(key=lambda x: (parse_date(x.get('DateCreated')) or datetime.date(1900,1,1), norm(x.get('NotesID'))))

def note_texts(oid):
    return [norm(n.get('Detail')) for n in notes_by_order.get(oid,[]) if norm(n.get('Detail'))]
def latest_notes(oid, n=3):
    arr=note_texts(oid)
    return ' || '.join(arr[-n:])[:1200]
def all_notes_lower(oid): return ' '.join(note_texts(oid)).lower()

# DNO tracker parse
with ZipFile(DNO_XLSX) as z:
    dno_data={}
    for tab in ['Pre-List','Domestic Orders','Service Replacement','Paid Quotes ','EVC Installations','Reporting']:
        dno_data[tab]=dict_rows_any(load_xlsx_sheet(z, tab))[1]
job_to_dno=defaultdict(list)
for tab,rows in dno_data.items():
    for d in rows:
        raw=norm(d.get('Job Number ') or d.get('Job Number') or '')
        for num in re.findall(r'\d{5,}', raw):
            job_to_dno[num].append((tab,d))

# Ranking table from docx for assumptions tab
ranking_rows=[]
try:
    doc=Document(str(DNO_DOCX))
    for table in doc.tables:
        for ri,row in enumerate(table.rows):
            cells=[c.text.strip() for c in row.cells]
            if ri>0 and len(cells)>=4 and cells[0].isdigit(): ranking_rows.append(cells[:4])
except Exception:
    pass

# ---------- classify notes ----------
def note_flags(oid):
    txt=all_notes_lower(oid)
    flags=[]
    if re.search(r'booked for|installation date|offer .*june|1st of june|confirmed scaff and installation', txt): flags.append('Booking/date mentioned in notes')
    # Be deliberately conservative with blockers: routine callbacks/chases should not block a June slot.
    if re.search(r'not ready|customer not ready|waiting (on|for) (a )?builder|awaiting (their )?builder|repointing|roof checked|local roofer|building works|scaffold(ing)? to be taken down|not ready just yet', txt): flags.append('Customer/roof/building readiness blocker')
    if re.search(r'finance.*unsigned|agreement.*unsigned|finance still showing as unsigned', txt): flags.append('Finance/document issue')
    if re.search(r'additional info requested|images required|meter / cut|meter/cut|cut.?out images|dno query|amendments required', txt): flags.append('DNO/additional information issue')
    if re.search(r'major works|required works|reinforcement|\bcoe\b|coes|chargeable|target date for a quote|export limit', txt) and not re.search(r'no works (are )?required|no works offer accepted|no export limits', txt): flags.append('DNO works/quote risk')
    if re.search(r'g99 approval received|approval received|fast track approval|no works offer accepted|full works accepted|accepted with no limits', txt): flags.append('DNO approval evidence in notes')
    if re.search(r'tried calling|left voicemail|emailed customer|chase|follow up|follow-up', txt): flags.append('Chase/customer contact required')
    return flags

def summarize_note_action(flags, status):
    if any('Customer/roof' in f for f in flags): return 'Confirm customer/roof/building readiness before confirming install date'
    if any('Finance' in f for f in flags): return 'Resolve finance/document issue before confirming install date'
    if any('DNO/additional' in f for f in flags): return 'Clear DNO/additional information request before confirming date'
    if any('DNO works' in f for f in flags): return 'Treat as DNO works/quote risk; confirm approval/COE before booking'
    if status in {'ValidationFailed'}: return 'Resolve validation failure, then confirm provisional slot'
    if status in {'InstallBooked','ScaffoldErectionRequired'}: return 'Red flag: verify why status has no installation date and correct source system'
    return 'Team to offer proposed provisional installation date subject to DNO/survey checks'

# ---------- enrich orders ----------
master=[]
for r in orders:
    oid=norm_id(r.get('OrderId'))
    pc=norm(r.get('CustomerPostcode') or r.get('CPostcode') or r.get('Postcode') or '')
    status=norm(r.get('Status'))
    sign_date=parse_date(r.get('SignDate'))
    build_date=parse_date(r.get('BuildDate'))
    flags=note_flags(oid)
    note_blob=' '.join(note_texts(oid))
    region=planning_region(pc, r.get('City'))
    mapped_dno, dno_source=dno_from_postcode(pc, note_blob)
    dno_matches=job_to_dno.get(oid,[])
    match_quality='Exact OrderId match' if dno_matches else 'No DNO tracker match'
    dno_status=''; dno_sent=None; dno_received=None; dno_notes=''; app_ref=''; prov_install=''
    service_blocker=''
    paid_quote=''
    for tab,d in dno_matches:
        if tab=='Domestic Orders':
            dno_status=norm(d.get('Application Status')) or dno_status
            dno_sent=parse_date(d.get('Date Application                Sent')) or dno_sent
            dno_received=parse_date(d.get('Date Application                Received')) or dno_received
            dno_notes=norm(d.get('Notes')) or dno_notes
            app_ref=norm(d.get('Application Reference')) or app_ref
            prov_install=norm(d.get('Provisional Installation Date')) or prov_install
        elif tab=='Service Replacement' and norm(d.get('Completed')).lower()!='yes':
            service_blocker=norm(d.get('Follow-Up Status'))
        elif tab=='Paid Quotes ':
            paid_quote=f"Quote paid status: {norm(d.get('Confirmation Received?'))} / {norm(d.get('Date Paid'))}"
        elif tab=='Pre-List':
            if not dno_status: dno_status='Pre-List: '+norm(d.get('Comments'))
    # notes can override evidence
    if not dno_received and any('DNO approval evidence' in f for f in flags):
        # keep tracker date empty but signal accepted evidence
        pass
    rule=DNO_RULES.get(mapped_dno, DNO_RULES['Unknown'])
    # Determine DNO readiness
    status_l=dno_status.lower()
    accepted = ('accepted' in status_l) or bool(dno_received) or any('DNO approval evidence' in f for f in flags)
    sent = bool(dno_sent) or ('sent' in status_l) or ('in review' in status_l)
    add_info = ('additional' in status_l) or any('DNO/additional' in f for f in flags)
    major = ('major' in status_l) or any('DNO works' in f for f in flags) or bool(service_blocker)
    g99_required = ('request required' in status_l) or ('pre-list' in status_l)
    wait_days=rule['typical']
    if accepted:
        expected_approval=dno_received or TODAY
        dno_bucket='Accepted / approval evidence'
        dno_conf='High'
    elif major:
        expected_approval=add_workdays(dno_sent or TODAY, rule['conservative'])
        dno_bucket='Major works / quote / technical risk'
        dno_conf='Low'
    elif add_info or g99_required or not sent:
        expected_approval=add_workdays(TODAY, rule['conservative'] if rule['rank']>=7 else max(rule['typical'],20))
        dno_bucket='Action needed before approval confidence'
        dno_conf='Low'
    else:
        expected_approval=add_workdays(dno_sent, wait_days)
        days_left=workdays_between(TODAY, expected_approval) if expected_approval else None
        dno_bucket='Submitted / waiting DNO'
        dno_conf='High' if expected_approval<=datetime.date(2026,6,10) and rule['rank']<=3 else ('Medium' if expected_approval<=datetime.date(2026,6,20) and rule['rank']<=6 else 'Low')
    # Survey assumption
    expected_survey=''
    if status=='AwaitingPhysicalSurveyBooking' and sign_date:
        expected_survey=add_workdays(sign_date,6)
    # Operational readiness
    blocker_flags=[]
    if status in {'InstallBooked','ScaffoldErectionRequired'}: blocker_flags.append('System status red flag')
    if status=='ValidationFailed': blocker_flags.append('Validation failed')
    if any('Customer/roof' in f for f in flags): blocker_flags.append('Customer/roof/building blocker')
    if any('Finance' in f for f in flags): blocker_flags.append('Finance/document blocker')
    if major: blocker_flags.append('DNO works/quote risk')
    if add_info: blocker_flags.append('DNO info issue')
    # Readiness bucket
    if status in {'InstallBooked','ScaffoldErectionRequired'}:
        readiness='Red flag / verify existing booking or scaffold position'
    elif blocker_flags:
        readiness='Conditional / action required before confirming'
    elif accepted:
        readiness='Ready to book into June'
    elif dno_conf=='High':
        readiness='Bookable June with DNO confidence'
    elif dno_conf=='Medium':
        readiness='Late-June provisional / monitor DNO'
    else:
        readiness='Late-June stretch / high risk'
    installer, installer_reason=installer_for(region, pc, mapped_dno, readiness)
    if installer=='Smartliving Energy': installer_risk='Difficult installer — controlled use only'
    elif installer=='CK Solar': installer_risk='Unreliable — avoid unless Dave approves'
    elif installer=='Solec Electrical Services' and region not in ['North East','Yorkshire / Humberside']:
        installer_risk='Solec national cluster only — do not send for isolated one-off'
    else:
        installer_risk=INSTALLERS.get(installer,{}).get('risk','')
    value=0.0
    try: value=float(norm(r.get('GrandTotal') or r.get('SaleSalePrice') or '0').replace(',',''))
    except Exception: value=0.0
    master.append({
        'OrderId':oid,'Customer':norm(r.get('CustomerFullName')),'Postcode':pc,'Outward':pc_outward(pc),'Area':pc_area(pc),'City':norm(r.get('City')),
        'Region':region,'Current Status':status,'Sales Channel':norm(r.get('SaleCompany')),'Sale Rep':norm(r.get('SaleRep')),
        'Sign Date':fmt_date(sign_date),'Build Date':fmt_date(build_date),'Panels':norm(r.get('#panel')),'Batteries':norm(r.get('#Battery')),
        'Tesla/Duracell': 'Tesla' if norm(r.get('Tesla')) else ('Duracell' if norm(r.get('Duracell')) else ''),'Order Value':value,
        'DNO':mapped_dno,'DNO Source':dno_source,'DNO Rank':rule['rank'],'DNO Risk':rule['risk'],'DNO Match':match_quality,
        'DNO Tracker Status':dno_status,'DNO Sent':fmt_date(dno_sent),'DNO Received':fmt_date(dno_received),'DNO Application Ref':app_ref,
        'Expected DNO Approval':fmt_date(expected_approval),'DNO Confidence':dno_conf,'DNO Bucket':dno_bucket,
        'Expected Survey By':fmt_date(expected_survey) if expected_survey else '',
        'Readiness Bucket':readiness,'Operational Flags':'; '.join(blocker_flags + flags[:4]),
        'Recommended Installer':installer,'Installer Rationale':installer_reason,'Installer Risk':installer_risk,
        'Internal Note Summary':latest_notes(oid,3),'Team Next Action':summarize_note_action(flags,status),
        '_score':0,'_expected_approval_date':expected_approval,'_value':value,
    })

# Scoring and scheduling
for row in master:
    score=0
    if row['DNO Confidence']=='High': score+=50
    elif row['DNO Confidence']=='Medium': score+=25
    else: score+=5
    if row['Readiness Bucket']=='Ready to book into June': score+=35
    elif row['Readiness Bucket']=='Bookable June with DNO confidence': score+=25
    elif 'Late-June provisional' in row['Readiness Bucket']: score+=10
    else: score-=15
    if 'red flag' in row['Readiness Bucket'].lower(): score-=35
    if 'blocker' in row['Operational Flags'].lower() or 'risk' in row['Operational Flags'].lower(): score-=20
    if row['DNO Rank']<=3: score+=15
    elif row['DNO Rank']<=6: score+=5
    else: score-=10
    if row['Order Value']>14000: score+=5
    row['_score']=score
    if score>=80: row['June Wave']='Wave 1: 1-5 Jun safest bookable'
    elif score>=55: row['June Wave']='Wave 2: 8-12 Jun high-confidence pending/accepted'
    elif score>=30: row['June Wave']='Wave 3: 15-19 Jun conditional/monitor'
    else: row['June Wave']='Wave 4: 22-30 Jun stretch/red-flag/action-led'

# June weekdays, split by wave
weekdays=[datetime.date(2026,6,d) for d in range(1,31) if datetime.date(2026,6,d).weekday()<5]
wave_dates={
    'Wave 1: 1-5 Jun safest bookable':[d for d in weekdays if d<=datetime.date(2026,6,5)],
    'Wave 2: 8-12 Jun high-confidence pending/accepted':[d for d in weekdays if datetime.date(2026,6,8)<=d<=datetime.date(2026,6,12)],
    'Wave 3: 15-19 Jun conditional/monitor':[d for d in weekdays if datetime.date(2026,6,15)<=d<=datetime.date(2026,6,19)],
    'Wave 4: 22-30 Jun stretch/red-flag/action-led':[d for d in weekdays if d>=datetime.date(2026,6,22)],
}
# Sort within each wave by region/installer clusters then score desc. Assign round-robin dates within allowed wave.
for wave, dates in wave_dates.items():
    rows=[r for r in master if r['June Wave']==wave]
    rows.sort(key=lambda r:(r['Recommended Installer'], r['Region'], -r['_score'], r['Outward']))
    counts=Counter()
    for i,r in enumerate(rows):
        # prefer date after expected approval unless wave 4 forced; choose first date >= expected approval + 1 workday if possible
        exp=r.get('_expected_approval_date')
        possible=dates[:]
        if exp and r['DNO Confidence']!='High':
            min_date=add_workdays(exp,1)
            later=[d for d in dates if d>=min_date]
            if later: possible=later
        # capacity spread: choose least used date in possible, preserving route clusters by installer/region
        chosen=min(possible, key=lambda d:(counts[(d,r['Recommended Installer'])], counts[d], d))
        counts[chosen]+=1; counts[(chosen,r['Recommended Installer'])]+=1
        r['Proposed Provisional Install Date']=fmt_date(chosen)
        r['Proposed Install Week']='W/C '+fmt_date(chosen-datetime.timedelta(days=chosen.weekday()))
        flags_text=(r.get('Operational Flags') or '').lower()
        if 'red flag' in r['Readiness Bucket'].lower() or 'customer/roof/building blocker' in flags_text or 'finance/document blocker' in flags_text or 'dno works/quote risk' in flags_text or 'dno info issue' in flags_text:
            r['Booking Confidence']='Low - provisional control date only'
        elif 'action required' in r['Readiness Bucket'].lower() and r['DNO Confidence']=='High':
            r['Booking Confidence']='Medium - clear action then offer date'
        elif r['DNO Confidence']=='Medium':
            r['Booking Confidence']='Medium - book with DNO monitoring'
        elif r['DNO Confidence']=='Low':
            r['Booking Confidence']='Low - provisional control date only'
        else:
            r['Booking Confidence']='High - offer date first'

# Run IDs: group by installer+week+region
run_counter=Counter()
for r in sorted(master, key=lambda x:(x['Recommended Installer'], x['Proposed Install Week'], x['Region'])):
    key=(r['Recommended Installer'], r['Proposed Install Week'], r['Region'])
    run_counter[key]+=1
run_ids={}
seq=1
for key,count in sorted(run_counter.items(), key=lambda kv:(kv[0][1], kv[0][0], kv[0][2])):
    run_ids[key]=f'RUN-{seq:02d}'
    seq+=1
for r in master:
    r['Installer Run ID']=run_ids[(r['Recommended Installer'], r['Proposed Install Week'], r['Region'])]
    run_size=run_counter[(r['Recommended Installer'], r['Proposed Install Week'], r['Region'])]
    r['Run Size']=run_size
    if r['Recommended Installer']=='Solec Electrical Services' and run_size<2 and r['Region'] not in ['North East','Yorkshire / Humberside']:
        r['Installer Risk'] += ' | Warning: insufficient cluster size for national Solec run'

# ---------- Workbook ----------
wb=Workbook()
# remove default
ws=wb.active; ws.title='Executive Summary'

blue='1F4E78'; navy='17365D'; green='92D050'; pale_green='E2F0D9'; amber='FCE4D6'; red='F4CCCC'; yellow='FFF2CC'; grey='D9EAD3'; lightblue='D9EAF7'
header_fill=PatternFill('solid', fgColor=blue)
sub_fill=PatternFill('solid', fgColor=lightblue)
thin=Side(style='thin', color='BFBFBF')
border=Border(left=thin,right=thin,top=thin,bottom=thin)

def style_sheet(ws, freeze='A2'):
    ws.freeze_panes=freeze
    ws.sheet_view.showGridLines=False
    for row in ws.iter_rows():
        for c in row:
            c.border=border
            c.alignment=Alignment(vertical='top', wrap_text=True)
    for c in ws[1]:
        c.fill=header_fill; c.font=Font(color='FFFFFF', bold=True); c.alignment=Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.auto_filter.ref=ws.dimensions

def write_rows(ws, headers, rows):
    ws.append(headers)
    for r in rows:
        ws.append([r.get(h,'') for h in headers])
    style_sheet(ws)
    widths={}
    for i,h in enumerate(headers,1):
        maxlen=len(str(h))
        for row in rows[:200]: maxlen=max(maxlen, len(str(row.get(h,''))) if row.get(h) is not None else 0)
        widths[i]=min(max(maxlen+2, 10), 45)
    for i,w in widths.items(): ws.column_dimensions[get_column_letter(i)].width=w
    if rows:
        tab=Table(displayName=re.sub('[^A-Za-z0-9_]','',ws.title)[:20]+'Tbl', ref=ws.dimensions)
        tab.tableStyleInfo=TableStyleInfo(name='TableStyleMedium2', showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
        ws.add_table(tab)

def apply_status_colours(ws, header_row=1):
    headers=[c.value for c in ws[header_row]]
    def col(name):
        return headers.index(name)+1 if name in headers else None
    conf_col=col('Booking Confidence'); wave_col=col('June Wave'); dno_col=col('DNO Confidence'); read_col=col('Readiness Bucket')
    for rr in range(2, ws.max_row+1):
        vals=' '.join(str(ws.cell(rr,c).value or '') for c in range(1,ws.max_column+1)).lower()
        fill=None
        if conf_col:
            v=str(ws.cell(rr,conf_col).value or '').lower()
            if v.startswith('high'): fill=PatternFill('solid', fgColor=pale_green)
            elif v.startswith('medium'): fill=PatternFill('solid', fgColor=yellow)
            elif v.startswith('low'): fill=PatternFill('solid', fgColor=amber)
        if 'red flag' in vals or 'validation failed' in vals or 'blocker' in vals:
            fill=PatternFill('solid', fgColor=red)
        if fill:
            for c in range(1, ws.max_column+1): ws.cell(rr,c).fill=fill

# Summary
ws=wb['Executive Summary']
summary=[
    ['Purpose','Turn 107 undated orders into a June booking control plan with DNO confidence, installer allocation and route/run logic.'],
    ['Planning bias','Earliest June dates go to accepted/high-confidence jobs; slower DNOs, blockers and red flags are pushed later as conditional dates.'],
    ['Operating rule','All orders have a June provisional control date; Low confidence rows should not be confirmed to customer until the action is cleared.'],
    ['CK Solar','Unreliable — avoided as default / last resort only.'],
    ['Smartliving','Difficult — controlled use, mainly Scotland only.'],
    ['Solec','Strategic cluster installer — national only when multiple jobs are booked consecutively in similar areas.'],
]
for r in summary: ws.append(r)
ws.append([])
# KPIs
kpis=[
    ['Total undated orders', len(master)],
    ['High booking confidence', sum(1 for r in master if str(r['Booking Confidence']).startswith('High'))],
    ['Medium booking confidence', sum(1 for r in master if str(r['Booking Confidence']).startswith('Medium'))],
    ['Low/control-date only', sum(1 for r in master if str(r['Booking Confidence']).startswith('Low'))],
    ['DNO exact tracker matches', sum(1 for r in master if r['DNO Match']=='Exact OrderId match')],
    ['No DNO tracker match', sum(1 for r in master if r['DNO Match']!='Exact OrderId match')],
    ['Red flag / action rows', sum(1 for r in master if 'red flag' in r['Readiness Bucket'].lower() or 'action required' in r['Readiness Bucket'].lower())],
]
ws.append(['Metric','Count'])
for r in kpis: ws.append(r)
ws.append([]); ws.append(['June wave','Orders'])
for w,c in Counter(r['June Wave'] for r in master).most_common(): ws.append([w,c])
ws.append([]); ws.append(['Installer','Orders'])
for k,c in Counter(r['Recommended Installer'] for r in master).most_common(): ws.append([k,c])
style_sheet(ws, freeze='A1')
ws.column_dimensions['A'].width=35; ws.column_dimensions['B'].width=90
for c in ws[1]: c.fill=header_fill; c.font=Font(color='FFFFFF',bold=True)

master_headers=['OrderId','Customer','Postcode','City','Region','Current Status','Sales Channel','Sign Date','Panels','Batteries','Tesla/Duracell','Order Value','DNO','DNO Rank','DNO Risk','DNO Match','DNO Tracker Status','DNO Sent','DNO Received','Expected DNO Approval','DNO Confidence','DNO Bucket','Expected Survey By','Readiness Bucket','June Wave','Proposed Install Week','Proposed Provisional Install Date','Booking Confidence','Recommended Installer','Installer Run ID','Run Size','Installer Rationale','Installer Risk','Operational Flags','Team Next Action','Internal Note Summary']
ws=wb.create_sheet('Master Order Planner')
write_rows(ws, master_headers, sorted(master, key=lambda r:(r['Proposed Provisional Install Date'], r['Recommended Installer'], r['Region'])))
apply_status_colours(ws)

# Include every order in the June Booking Plan because Dave wants a June control date for every undated order.
# Confidence/action columns tell the team which dates can be offered immediately vs held as provisional control dates.
june_rows=master
ws=wb.create_sheet('June Booking Plan')
booking_headers=['Proposed Provisional Install Date','Proposed Install Week','Installer Run ID','Recommended Installer','Run Size','Region','OrderId','Customer','Postcode','Current Status','DNO','DNO Tracker Status','Expected DNO Approval','DNO Confidence','Booking Confidence','Team Next Action','Installer Rationale']
write_rows(ws, booking_headers, sorted(june_rows, key=lambda r:(r['Proposed Provisional Install Date'], r['Recommended Installer'], r['Region'])))
apply_status_colours(ws)

exc_rows=[r for r in master if str(r['Booking Confidence']).startswith('Low') or 'red flag' in r['Readiness Bucket'].lower() or 'blocker' in r['Operational Flags'].lower() or 'No DNO' in r['DNO Match']]
ws=wb.create_sheet('Red Flags & Actions')
exc_headers=['Proposed Provisional Install Date','OrderId','Customer','Postcode','Current Status','Readiness Bucket','Operational Flags','DNO','DNO Tracker Status','DNO Confidence','Recommended Installer','Team Next Action','Internal Note Summary']
write_rows(ws, exc_headers, sorted(exc_rows, key=lambda r:(r['Proposed Provisional Install Date'], r['Current Status'], r['OrderId'])))
apply_status_colours(ws)

chase_rows=[r for r in master if r['DNO Confidence']!='High' or r['DNO Match']!='Exact OrderId match' or 'Submitted' in r['DNO Bucket'] or 'Action needed' in r['DNO Bucket']]
ws=wb.create_sheet('DNO Chase List')
chase_headers=['OrderId','Customer','Postcode','DNO','DNO Rank','DNO Risk','DNO Match','DNO Tracker Status','DNO Sent','Expected DNO Approval','DNO Confidence','DNO Bucket','Proposed Provisional Install Date','Team Next Action','Internal Note Summary']
write_rows(ws, chase_headers, sorted(chase_rows, key=lambda r:(r['DNO Confidence'], r['Expected DNO Approval'], r['DNO Rank'])))
apply_status_colours(ws)

# Installer runs aggregated
run_rows=[]
for key, size in run_counter.items():
    inst, week, region=key
    rows=[r for r in master if (r['Recommended Installer'],r['Proposed Install Week'],r['Region'])==key]
    dates=sorted(set(r['Proposed Provisional Install Date'] for r in rows))
    run_rows.append({
        'Installer Run ID':run_ids[key],'Installer':inst,'Week':week,'Region':region,'Orders':size,
        'Dates':', '.join(dates),'Order IDs':', '.join(r['OrderId'] for r in sorted(rows,key=lambda x:x['Proposed Provisional Install Date'])),
        'DNO Confidence Mix':'; '.join(f'{k}:{v}' for k,v in Counter(r['DNO Confidence'] for r in rows).items()),
        'Booking Confidence Mix':'; '.join(f'{k}:{v}' for k,v in Counter(r['Booking Confidence'] for r in rows).items()),
        'Route / Executive Note': ('Good clustered run' if size>=2 else 'Single job — check travel efficiency before confirming') + ('; Solec national rule applies' if inst=='Solec Electrical Services' else '')
    })
ws=wb.create_sheet('Installer Runs')
run_headers=['Installer Run ID','Installer','Week','Region','Orders','Dates','Order IDs','DNO Confidence Mix','Booking Confidence Mix','Route / Executive Note']
write_rows(ws, run_headers, sorted(run_rows, key=lambda r:(r['Week'], r['Installer'], r['Region'])))
apply_status_colours(ws)

# Source notes compact
note_rows=[]
for r in master:
    oid=r['OrderId']; arr=notes_by_order.get(oid,[])
    note_rows.append({'OrderId':oid,'Customer':r['Customer'],'Note Count':len(arr),'Latest Notes':latest_notes(oid,5)})
ws=wb.create_sheet('Order Notes Summary')
write_rows(ws, ['OrderId','Customer','Note Count','Latest Notes'], sorted(note_rows, key=lambda r:r['OrderId']))

# Installer network
inst_out=[]
for name,meta in INSTALLERS.items():
    inst_out.append({'Installer':name,'Coverage / Use':', '.join(meta['region']),'Status':meta['status'],'Planner Risk / Rule':meta['risk'],'Allocated Orders':sum(1 for r in master if r['Recommended Installer']==name)})
ws=wb.create_sheet('Installer Network')
write_rows(ws, ['Installer','Coverage / Use','Status','Planner Risk / Rule','Allocated Orders'], inst_out)

# DNO assumptions
ass=[]
for d,rule in sorted(DNO_RULES.items(), key=lambda kv: kv[1]['rank']):
    if d=='Unknown': continue
    ass.append({'Rank':rule['rank'],'DNO / IDNO':d,'Typical working days used':rule['typical'],'Conservative working days':rule['conservative'],'Risk':rule['risk'],'Planner note':rule['note']})
ws=wb.create_sheet('DNO Assumptions')
write_rows(ws, ['Rank','DNO / IDNO','Typical working days used','Conservative working days','Risk','Planner note'], ass)

# Assumptions and source control
ass2=[
    {'Rule Area':'Status','Rule':'InstallBooked / scaffold statuses with no date are red flags; given late June control dates until verified.'},
    {'Rule Area':'Survey','Rule':'AwaitingPhysicalSurveyBooking assumed surveyed within six working days from signed date.'},
    {'Rule Area':'DNO','Rule':'Accepted/approval evidence gets earliest dates; submitted rows use DNO average wait; slow/high-risk DNOs move later.'},
    {'Rule Area':'DNO','Rule':'Postcode-to-DNO is planning intelligence; team should verify against actual application/ref before customer confirmation.'},
    {'Rule Area':'Installer','Rule':'Solec national work only when clustered in same region/week; CK avoided; Smartliving controlled use only.'},
    {'Rule Area':'Output','Rule':'Every order has a June provisional control date, but Low confidence rows are not customer-confirmation ready.'},
]
ws=wb.create_sheet('Planner Assumptions')
write_rows(ws, ['Rule Area','Rule'], ass2)

# Apply workbook-wide polish
for ws in wb.worksheets:
    ws.sheet_view.showGridLines=False
    # row heights
    for row in range(1, min(ws.max_row,300)+1): ws.row_dimensions[row].height=30 if row>1 else 36

# Save
OUT.parent.mkdir(parents=True, exist_ok=True)
wb.save(OUT)
print(json.dumps({
    'output':str(OUT), 'orders':len(master), 'high':sum(1 for r in master if str(r['Booking Confidence']).startswith('High')), 'medium':sum(1 for r in master if str(r['Booking Confidence']).startswith('Medium')), 'low':sum(1 for r in master if str(r['Booking Confidence']).startswith('Low')), 'sheets':wb.sheetnames
}, indent=2))
